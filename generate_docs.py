#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为软著申请生成源代码和使用说明书Word文档"""

import os
import glob
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = r"c:/Users/iuuuu/AndroidStudioProjects/Studyapp"
SRC_DIR = os.path.join(PROJECT_DIR, "app/src/main/java/com/example/studyapp")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "软著申请文档")
APP_NAME = "StudyApp学习平台"
APP_NAME_SHORT = "StudyApp"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(doc, text, bold=False, font_size=12, alignment=None, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def add_img_placeholder(doc, description, img_index):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ 截图{img_index} ]')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'图{img_index}：{description}')
    run2.font.size = Pt(10)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('  • ' + text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(2)


def add_step(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('  ' + text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(3)


# ============================================================
# 1. 源代码处理
# ============================================================

def remove_kotlin_comments(text):
    """移除Kotlin代码中的注释"""
    result = []
    i = 0
    while i < len(text):
        if text[i] == '"':
            result.append('"')
            i += 1
            while i < len(text) and text[i] != '"':
                if text[i] == '\\':
                    result.append(text[i])
                    i += 1
                if i < len(text):
                    result.append(text[i])
                    i += 1
            if i < len(text):
                result.append('"')
                i += 1
        elif text[i:i+2] == '/*':
            depth = 1
            i += 2
            while i < len(text) and depth > 0:
                if text[i:i+2] == '/*':
                    depth += 1
                    i += 2
                elif text[i:i+2] == '*/':
                    depth -= 1
                    i += 2
                else:
                    i += 1
        elif text[i:i+2] == '//':
            i += 2
            while i < len(text) and text[i] != '\n':
                i += 1
        else:
            result.append(text[i])
            i += 1
    return ''.join(result)


def get_source_files():
    files = []
    for root, dirs, filenames in os.walk(SRC_DIR):
        for fn in sorted(filenames):
            if fn.endswith('.kt'):
                full_path = os.path.join(root, fn)
                rel_path = os.path.relpath(full_path, SRC_DIR)
                files.append((rel_path, full_path))
    return sorted(files, key=lambda x: x[0])


def generate_source_code():
    print('=' * 60)
    print('正在生成源代码文档...')
    print('=' * 60)

    files = get_source_files()
    print(f'找到 {len(files)} 个Kotlin源文件')

    all_lines = []
    for rel_path, full_path in files:
        with open(full_path, 'r', encoding='utf-8') as f:
            content = f.read()
        clean = remove_kotlin_comments(content)
        file_lines = clean.split('\n')
        all_lines.append(f'// ===== {rel_path} =====')
        all_lines.extend(file_lines)
        all_lines.append('')

    total_lines = len(all_lines)
    print(f'总行数（清理注释后）: {total_lines}')

    first_part = all_lines[:1500]
    last_part = all_lines[-1500:] if total_lines > 1500 else []

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(7)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Title page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{APP_NAME}')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('源 代 码')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n（共{total_lines}行，取前1500行和后1500行）')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def write_code_table(doc, lines, start_line_num):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for cell in table.columns[0].cells:
            cell.width = Cm(1.5)
        for cell in table.columns[1].cells:
            cell.width = Cm(14.5)

        for i, line in enumerate(lines):
            if i > 0:
                table.add_row()
            row_cells = table.rows[-1].cells
            row_cells[0].text = str(start_line_num + i)
            for para in row_cells[0].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in para.runs:
                    r.font.size = Pt(7)
                    r.font.name = 'Courier New'
            row_cells[1].text = line
            for para in row_cells[1].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(7)
                    r.font.name = 'Courier New'

    doc.add_page_break()
    add_heading_styled(doc, '一、源代码（前半部分，第1-1500行）', level=1)
    write_code_table(doc, first_part, 1)

    doc.add_page_break()
    add_heading_styled(doc, f'二、源代码（后半部分，第{total_lines-1499}-{total_lines}行）', level=1)
    if last_part:
        write_code_table(doc, last_part, total_lines - 1500 + 1)

    output_path = os.path.join(OUTPUT_DIR, f'{APP_NAME}_源代码.docx')
    doc.save(output_path)
    print(f'[OK] 源代码文档已生成: {output_path}')
    return output_path


# ============================================================
# 2. 使用说明书
# ============================================================

def generate_user_manual():
    print('\n' + '=' * 60)
    print('正在生成使用说明书文档...')
    print('=' * 60)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ====== Cover ======
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{APP_NAME}')
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('使 用 说 明 书')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for _ in range(4):
        doc.add_paragraph()

    for label, value in [
        ('软件名称', APP_NAME),
        ('版本号', 'V1.0'),
        ('开发单位', '个人开发者'),
        ('文档版本', 'V1.0'),
        ('编制日期', '2026年5月'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ====== TOC ======
    add_heading_styled(doc, '目 录', level=1)
    toc = [
        '1. 引言',
        '   1.1 编写目的',
        '   1.2 软件概述',
        '2. 运行环境',
        '   2.1 硬件环境',
        '   2.2 软件环境',
        '3. 软件安装与启动',
        '   3.1 安装步骤',
        '   3.2 启动程序',
        '4. 登录与注册',
        '   4.1 用户登录',
        '   4.2 用户注册',
        '   4.3 忘记密码',
        '5. 学生端功能',
        '   5.1 首页概览',
        '   5.2 我的课程',
        '   5.3 选课中心',
        '   5.4 消息中心',
        '   5.5 视频播放',
        '6. 教师端功能',
        '   6.1 首页概览',
        '   6.2 上传课程',
        '   6.3 课程管理',
        '   6.4 数据统计',
        '7. 管理员端功能',
        '   7.1 仪表盘',
        '   7.2 用户管理',
        '   7.3 活跃度排行',
        '   7.4 学习统计',
        '   7.5 课程掌握度',
        '8. 常见问题与故障排除',
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(20)

    doc.add_page_break()

    # ====== 1. Intro ======
    add_heading_styled(doc, '1. 引言', level=1)
    add_heading_styled(doc, '1.1 编写目的', level=2)
    add_para(doc,
        '本文档为《StudyApp学习平台》软件的使用说明书，旨在帮助用户快速熟悉和掌握本软件的各项功能操作。'
        '本说明书详细介绍了软件的安装部署、登录注册、各角色功能模块的使用方法，'
        '配合系统截图和文字说明，使用户能够直观地了解每个操作步骤。'
        '本软件是一款面向移动端的在线学习管理平台，支持学生、教师和管理员三种角色，'
        '提供课程学习、视频播放、数据统计等核心功能。')

    add_heading_styled(doc, '1.2 软件概述', level=2)
    add_para(doc,
        'StudyApp学习平台是一款基于Android平台的在线学习管理系统。'
        '系统采用前后端分离架构，前端使用Kotlin语言开发，后端基于Node.js和MySQL数据库。'
        '学生可以在线选择课程、观看教学视频、参与测验；'
        '教师可以上传和管理课程内容、查看学生学习统计数据；'
        '管理员可以管理用户账号、监控系统运行状态、查看多维度的学习数据分析报表。')

    doc.add_page_break()

    # ====== 2. Environment ======
    add_heading_styled(doc, '2. 运行环境', level=1)
    add_heading_styled(doc, '2.1 硬件环境', level=2)
    add_para(doc, '客户端（Android设备）：')
    for item in [
        '处理器：ARM架构，1.5GHz及上',
        '内存：2GB及上',
        '存储空间：100MB及上可用空间',
        '屏幕分辨率：1280×720及上',
        '网络：支持Wi-Fi或移动数据网络',
    ]:
        add_bullet(doc, item)

    add_para(doc, '')
    add_para(doc, '服务器端：')
    for item in [
        'CPU：2核及上',
        '内存：4GB及上',
        '硬盘：40GB及上',
        '操作系统：Linux CentOS 7.x 或 Ubuntu 20.x',
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, '2.2 软件环境', level=2)
    for item in [
        '客户端：Android 7.0 (API 24) 及上版本',
        '服务端：Node.js 16.x + MySQL 8.x',
        '开发框架：Android SDK + ExoPlayer + Retrofit + Coil',
        '后端依赖：Express.js、mysql2、cors、body-parser',
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # ====== 3. Install ======
    add_heading_styled(doc, '3. 软件安装与启动', level=1)
    add_heading_styled(doc, '3.1 安装步骤', level=2)
    add_para(doc, '本软件以APK安装包形式分发。用户可通过以下方式安装：')
    for step in [
        '（1）方式一：将APK文件传输至Android设备，使用文件管理器找到APK文件，点击即可进入安装向导，按照屏幕提示完成安装。',
        '（2）方式二：通过USB数据线连接Android设备和开发电脑，在终端中使用 adb install 命令进行安装。',
        '（3）安装完成后，在Android设备的应用列表中找到「StudyApp学习平台」图标，点击即可启动。',
    ]:
        add_step(doc, step)

    add_img_placeholder(doc, 'APK安装界面截图', 1)

    add_heading_styled(doc, '3.2 启动程序', level=2)
    add_para(doc,
        '安装完成后，点击应用图标启动程序。系统将显示欢迎界面，随后进入登录页面。'
        '首次使用时，用户需要注册账号或使用已有账号登录。')

    add_img_placeholder(doc, '应用启动界面截图', 2)

    doc.add_page_break()

    # ====== 4. Login & Register ======
    add_heading_styled(doc, '4. 登录与注册', level=1)
    add_heading_styled(doc, '4.1 用户登录', level=2)
    add_para(doc,
        '打开应用后，进入登录界面。登录界面包含用户名输入框、密码输入框、'
        '「记住密码」复选框、登录按钮、注册链接和忘记密码链接。操作步骤如下：')
    for step in [
        '（1）在用户名输入框中输入已注册的用户名（如教师账号 teacher001）；',
        '（2）在密码输入框中输入对应的密码（默认密码 123456）；',
        '（3）勾选「记住密码」复选框可保存登录凭据，方便下次快速登录；',
        '（4）点击「登录」按钮，系统验证身份后自动跳转至对应角色主界面。',
    ]:
        add_step(doc, step)

    add_img_placeholder(doc, '登录界面截图', 3)

    add_para(doc,
        '系统支持学生（student）、教师（teacher）和管理员（admin）三种角色，'
        '不同角色登录后进入不同的功能界面。'
        '登录成功后，系统会保存登录状态，退出应用后重新打开无需再次登录。')

    add_heading_styled(doc, '4.2 用户注册', level=2)
    add_para(doc,
        '若用户尚未拥有账号，可点击登录界面的「注册」链接进入注册页面。注册页面包含以下信息填写项：')
    for item in [
        '用户名：2-20位字符，支持字母、数字、下划线和中文',
        '密码：至少6位',
        '确认密码：再次输入密码以确认',
        '真实姓名：填写用户真实姓名',
        '生日：通过日期选择器设置出生日期',
        '角色：选择学生、教师或管理员（管理员需填写验证码）',
        '密保问题：从预设问题中选择一个',
        '密保答案：填写密保问题的答案（用于找回密码）',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '注册界面截图', 4)

    add_heading_styled(doc, '4.3 忘记密码', level=2)
    add_para(doc,
        '如果用户忘记密码，可点击登录界面的「忘记密码」链接。'
        '系统将弹出忘记密码对话框，操作流程如下：')
    for step in [
        '（1）输入用户名，点击「查询」按钮获取密保问题；',
        '（2）系统显示预设的密保问题，用户需要输入对应的答案；',
        '（3）同时需要验证真实姓名和生日信息；',
        '（4）验证通过后，输入新密码并确认；',
        '（5）点击「重置密码」完成密码修改，使用新密码登录。',
    ]:
        add_step(doc, step)

    add_img_placeholder(doc, '忘记密码对话框截图', 5)

    doc.add_page_break()

    # ====== 5. Student ======
    add_heading_styled(doc, '5. 学生端功能', level=1)
    add_para(doc,
        '学生登录成功后进入学生端主界面。主界面采用侧边栏导航设计，'
        '包含首页、我的课程、选课中心和消息四个核心功能模块。')

    add_img_placeholder(doc, '学生端主界面截图', 6)

    add_heading_styled(doc, '5.1 首页概览', level=2)
    add_para(doc, '学生端首页展示了学习相关的概览信息，包括：')
    for item in [
        '欢迎区域：显示当前登录用户的用户名和个性化问候',
        '学习日历：以日历形式展示每日学习记录，已学习的日期有标记',
        '学习进度概览：展示当前已选课程的数量和学习进度统计',
        '快捷入口：提供课程、选课等功能的快速跳转按钮',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '学生端首页截图', 7)

    add_heading_styled(doc, '5.2 我的课程', level=2)
    add_para(doc,
        '「我的课程」模块展示学生已选的所有课程列表。每门课程以卡片形式展示，'
        '包含课程封面图、课程名称、授课教师、课程简介和课程学分信息。'
        '点击课程卡片可进入课程详情页面，查看课程视频列表和学习资料。')
    add_para(doc,
        '课程详情页显示该课程的所有教学视频列表，学生可以点击视频进行在线播放学习。'
        '系统会自动记录学习进度，下次进入时可从上次中断位置继续观看。')

    add_img_placeholder(doc, '我的课程列表截图', 8)

    add_heading_styled(doc, '5.3 选课中心', level=2)
    add_para(doc,
        '「选课中心」模块展示系统所有可选的课程列表。学生可以浏览全部课程，'
        '查看课程详细信息，并点击「选课」按钮将课程添加到自己的课程列表中。'
        '选课后，该课程会出现在「我的课程」模块中。')

    add_img_placeholder(doc, '选课中心界面截图', 9)

    add_heading_styled(doc, '5.4 消息中心', level=2)
    add_para(doc,
        '「消息中心」模块提供系统消息通知功能。学生可以在此查看系统推送的通知消息，'
        '包括课程更新通知、学习提醒等信息。')

    add_img_placeholder(doc, '消息中心界面截图', 10)

    add_heading_styled(doc, '5.5 视频播放', level=2)
    add_para(doc,
        '视频播放器是学生端核心功能之一，支持在线播放教学视频。播放器具有以下功能：')
    for item in [
        '播放/暂停控制',
        '进度条拖动，支持跳转到任意播放位置',
        '全屏模式切换',
        '倍速播放（支持0.5x、1.0x、1.5x、2.0x）',
        '播放进度自动保存，下次继续播放',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '视频播放界面截图', 11)

    doc.add_page_break()

    # ====== 6. Teacher ======
    add_heading_styled(doc, '6. 教师端功能', level=1)
    add_para(doc,
        '教师登录成功后进入教师端主界面。教师端采用侧边栏导航设计，'
        '包含首页、上传课程、我的课程和数据统计四个核心模块。')

    add_img_placeholder(doc, '教师端主界面截图', 12)

    add_heading_styled(doc, '6.1 首页概览', level=2)
    add_para(doc,
        '教师端首页展示教学相关的概览信息，包括个人信息展示和快捷功能入口。'
        '教师可以快速了解当前的教学概况。')

    add_img_placeholder(doc, '教师端首页截图', 13)

    add_heading_styled(doc, '6.2 上传课程', level=2)
    add_para(doc,
        '「上传课程」模块提供给教师创建和发布新课程的功能。上传课程的表单包含以下填写项：')
    for item in [
        '课程名称：输入课程标题',
        '课程简介：编写课程描述和教学目标',
        '课程学分：设置课程的学分值',
        '课程封面：支持选择图片作为课程封面，同时支持AI自动生成课程封面图',
        '上传视频：选择本地视频文件上传至服务器作为课程教学视频',
    ]:
        add_bullet(doc, item)

    add_para(doc,
        'AI封面生成功能使用先进的AI模型，根据课程名称自动生成匹配的课程封面图片，'
        '简化了教师制作课程封面的流程。视频上传支持MP4格式，系统会自动处理视频转码。')

    add_img_placeholder(doc, '上传课程界面截图', 14)

    add_heading_styled(doc, '6.3 课程管理', level=2)
    add_para(doc,
        '「我的课程」模块展示该教师已创建的所有课程列表。教师可以查看每门课程的详细信息，'
        '包括选课学生人数等统计数据。点击课程可进入课程详情页，管理课程视频内容。')

    add_img_placeholder(doc, '教师课程管理界面截图', 15)

    add_heading_styled(doc, '6.4 数据统计', level=2)
    add_para(doc,
        '「数据统计」模块为教师提供多维度的教学数据分析，包括：')
    for item in [
        '选课人数统计：展示每门课程的选课学生总数',
        '平均观看时长：统计学生对教学视频的平均观看时间',
        '学习完成率：展示课程学习进度的完成比例分布',
        '点击率统计：统计课程内容的访问和互动数据',
        '测验成绩分析：展示学生在线测验的成绩分布情况',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '教师数据统计界面截图', 16)

    doc.add_page_break()

    # ====== 7. Admin ======
    add_heading_styled(doc, '7. 管理员端功能', level=1)
    add_para(doc,
        '管理员登录成功后进入管理员端主界面。管理员端采用侧边栏导航设计，'
        '包含仪表盘、用户管理、活跃度排行、学习统计和课程掌握度五个核心模块。')

    add_img_placeholder(doc, '管理员端主界面截图', 17)

    add_heading_styled(doc, '7.1 仪表盘', level=2)
    add_para(doc,
        '仪表盘是管理员登录后的默认页面，展示系统的整体运行概览。仪表盘包含以下信息模块：')
    for item in [
        '在线用户数：实时显示当前在线用户数量，帮助管理员了解系统活跃情况',
        '活跃度排行：展示学习活跃度最高的前10名用户排名列表',
        '学习统计：汇总展示系统的总体学习数据，包括总学习时长、学习人次等',
        '课程掌握度：以可视化方式展示各门课程的学生整体掌握水平',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '管理员仪表盘截图', 18)

    add_heading_styled(doc, '7.2 用户管理', level=2)
    add_para(doc,
        '「用户管理」模块提供系统用户的全面管理功能。管理员可以：')
    for item in [
        '查看所有注册用户的列表，包括用户名、角色、注册时间等信息',
        '搜索和筛选用户，支持按用户名和角色进行快速查找',
        '编辑用户信息，修改用户角色或重置密码',
        '删除用户账号，同时自动清理该用户相关的选课记录和学习记录',
        '添加新用户，直接为特定用户创建账号',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '用户管理界面截图', 19)

    add_heading_styled(doc, '7.3 活跃度排行', level=2)
    add_para(doc,
        '「活跃度排行」模块展示系统用户的活跃度排名。该排行基于用户的学习记录数据，'
        '包括学习时长、登录频率、课程参与度等多个维度计算得出。'
        '排行榜支持按日、周、月等时间维度进行筛选查看。')

    add_img_placeholder(doc, '活跃度排行界面截图', 20)

    add_heading_styled(doc, '7.4 学习统计', level=2)
    add_para(doc,
        '「学习统计」模块提供系统级的详细学习数据分析。管理员可以查看：')
    for item in [
        '总学习时长统计：显示所有学生的累计学习时间',
        '各课程学习热度：分析不同课程的受欢迎程度和学习活跃度',
        '学习趋势图：以图表形式展示一段时间内的学习活跃度变化趋势',
        '学生参与度分析：统计学生参与课程学习、测验等活动的比例',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '学习统计界面截图', 21)

    add_heading_styled(doc, '7.5 课程掌握度', level=2)
    add_para(doc,
        '「课程掌握度」模块以可视化的方式展示每门课程的掌握度分析。'
        '系统通过分析学生的视频观看进度、测验成绩、学习时长等数据，'
        '综合评估学生对各门课程的掌握程度。管理员可以直观地了解：')
    for item in [
        '各课程的整体掌握度评分',
        '不同学生群体在各课程中的表现差异',
        '需要重点关注的课程或学生群体',
    ]:
        add_bullet(doc, item)

    add_img_placeholder(doc, '课程掌握度界面截图', 22)

    doc.add_page_break()

    # ====== 8. FAQ ======
    add_heading_styled(doc, '8. 常见问题与故障排除', level=1)

    faqs = [
        ('Q1: 无法登录怎么办？',
         '请检查用户名和密码是否正确。如果忘记密码，可以使用「忘记密码」功能通过密保问题重置密码。如仍无法登录，请联系系统管理员。'),
        ('Q2: 视频播放卡顿或无法播放？',
         '请检查网络连接是否正常。可以尝试切换Wi-Fi和移动数据网络。如果问题持续，可能是视频文件正在转码处理中，请稍后再试。'),
        ('Q3: 选课后课程不显示？',
         '选课成功后课程会出现在「我的课程」列表中。如果未显示，请下拉刷新页面。如仍无法显示，请重新登录后查看。'),
        ('Q4: 如何修改个人信息？',
         '目前个人信息修改功能在个人资料页面。您可以修改头像等基本信息。如需修改用户名或角色，请联系管理员。'),
        ('Q5: APK安装失败？',
         '请在设备设置中开启「允许安装未知来源应用」选项。建议通过文件管理器找到APK文件后点击安装，避免直接在微信或QQ中打开安装。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        run = p.add_run(a)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(8)

    # ====== Appendix ======
    doc.add_page_break()
    add_heading_styled(doc, '附录：系统界面索引', level=1)
    add_para(doc,
        '本文档共包含以下系统界面截图索引，方便用户快速定位各个功能模块的操作说明：')

    for img_num, desc in [
        ('图1', 'APK安装界面'),
        ('图2', '应用启动界面'),
        ('图3', '登录界面'),
        ('图4', '注册界面'),
        ('图5', '忘记密码对话框'),
        ('图6', '学生端主界面'),
        ('图7', '学生端首页'),
        ('图8', '我的课程列表'),
        ('图9', '选课中心界面'),
        ('图10', '消息中心界面'),
        ('图11', '视频播放界面'),
        ('图12', '教师端主界面'),
        ('图13', '教师端首页'),
        ('图14', '上传课程界面'),
        ('图15', '教师课程管理界面'),
        ('图16', '教师数据统计界面'),
        ('图17', '管理员端主界面'),
        ('图18', '管理员仪表盘'),
        ('图19', '用户管理界面'),
        ('图20', '活跃度排行界面'),
        ('图21', '学习统计界面'),
        ('图22', '课程掌握度界面'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f'  {img_num}：{desc}')
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = Pt(2)

    output_path = os.path.join(OUTPUT_DIR, f'{APP_NAME}_使用说明书.docx')
    doc.save(output_path)
    print(f'[OK] 使用说明书文档已生成: {output_path}')
    return output_path


# ============================================================

if __name__ == '__main__':
    print('=' * 60)
    print(f'  {APP_NAME} - 软著申请文档生成工具')
    print('=' * 60)

    src_path = generate_source_code()
    manual_path = generate_user_manual()

    print('\n' + '=' * 60)
    print('[OK] 所有文档生成完毕！')
    print('=' * 60)
    print(f'📄 源代码文档: {src_path}')
    print(f'📄 使用说明书: {manual_path}')
    print()
    print('📍 请在使用说明书中插入以下位置的截图：')
    print('    共 22 张截图，详见文档中的[截图N]标记')
    print('\n💡 打开Word文档后，将[截图N]标记替换为实际截图即可')